# pip install python-docx

from docx.api import Document
 

import os
import re

def doc2docx(path):
    import win32com.client as win32
    from win32com.client import constants
    import os
    import re
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(path)
    doc.Activate()
    new_file_abs_path = os.path.abspath(path)
    new_file_abs_path = re.sub(r'\.\w+$', '.docx', new_file_abs_path)
    word.ActiveDocument.SaveAs(new_file_abs_path, FileFormat=constants.wdFormatXMLDocument)
    doc.Close()

docfiles = [f for f in os.listdir('.') if os.path.isfile(f) and f.endswith('.doc')]
# basepath = 'C:\\Users\\dieze\\OneDrive - Brock University\\TA\\5P12\\2021F\\PEER REVIEW\\all\\'
basepath = 'C:\\Users\\dieze\\OneDrive - Brock University\\TA\\5P11\\2021F\\PEER REVIEW\\all'

for docfile in docfiles:
    doc2docx(os.path.join(basepath,docfile))

from docx.api import Document
import pandas as pd

doxfiles = [f for f in os.listdir('.') if os.path.isfile(f) and f.endswith('.docx')]

results =pd.DataFrame()
# document = Document('C:\\Users\\dieze\\OneDrive - Brock University\\TA\\5P12\\2021F\\PEER REVIEW\\all\\., Kirti(ks20pl)_Peer+Evaluation+Form.docx')
for dox in doxfiles:
    document = Document(os.path.join(basepath,dox))
    tables = document.tables
    names = [p.text.split('Name')[-1].replace('_','').strip() for p in document.paragraphs if 'Team Member' in p.text]
    names = [n for n in names if n]

    # docx table to pandas
    for i, table in enumerate(document.tables[:len(names)]):
        df = pd.DataFrame()
        for row in table.rows:
            text = [cell.text for cell in row.cells]
            df = df.append([text], ignore_index=True)
        df.columns = df.loc[0]
        df = df.drop(0)
        
        df.set_index('FACTOR', inplace=True)
        for f in df.index:
            for score in ['A', 'B', 'C', 'D']:
                if str(df.loc[f, score]).strip():
                    df.loc[f, 'mark'] = score
                    break
        if 'mark' not in df:
            df['mark'] = 'A'
        result = df['mark'].to_frame().T
        result['name'] = re.sub(r'\W+',' ',names[i]).strip()
        results = pd.concat([results, result])

results.to_csv('evaluations.csv', index=False)
